"""Simple Markdown -> DOCX converter using python-docx

Usage:
  pip install python-docx
  python scripts/generate_docx_from_md.py ../FINAL_REPORT.md ../FINAL_REPORT.docx

This converter maps Markdown headings (#, ##, ###) to Word heading styles and writes paragraphs.
It does not embed images. It outputs a .docx suitable for submission.
"""
import sys
from docx import Document
from docx.shared import Pt

HEADING_MAP = {
    1: 'Title',
    2: 'Heading 1',
    3: 'Heading 2',
    4: 'Heading 3'
}

def md_to_docx(md_path, docx_path):
    doc = Document()
    with open(md_path, 'r', encoding='utf-8') as f:
        for raw_line in f:
            line = raw_line.rstrip('\n')
            if not line:
                doc.add_paragraph('')
                continue
            # Heading detection
            if line.startswith('#'):
                # count hashes
                hashes = 0
                for ch in line:
                    if ch == '#':
                        hashes += 1
                    else:
                        break
                text = line[hashes:].strip()
                style = HEADING_MAP.get(hashes, 'Heading 4')
                doc.add_paragraph(text, style)
            else:
                # regular paragraph
                p = doc.add_paragraph(line)
                # simple inline formatting for bold/italic is not implemented
    doc.save(docx_path)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: python generate_docx_from_md.py input.md output.docx')
        sys.exit(1)
    md_to_docx(sys.argv[1], sys.argv[2])
    print('Saved', sys.argv[2])
